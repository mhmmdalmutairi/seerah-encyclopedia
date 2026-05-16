#!/usr/bin/env python3
"""Build the 3-page Arabic strategic brief as a Word document.
Content mirrors markaz/index.html (the HTML deck). Strict 3-page budget,
consistent typography, RTL tables, color-coded callouts.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ARABIC_FONT = 'IBM Plex Sans Arabic'

PRIMARY = (10, 77, 104)      # #0A4D68
ACCENT = (195, 155, 92)      # #C39B5C
TEXT = (31, 41, 55)          # #1F2937
MUTED = (107, 114, 128)      # #6B7280


def _arabic_font(run, size, bold=False, color=TEXT):
    run.font.name = ARABIC_FONT
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:cs'), ARABIC_FONT)
    rFonts.set(qn('w:hAnsi'), ARABIC_FONT)
    rFonts.set(qn('w:ascii'), ARABIC_FONT)
    run.font.size = Pt(size)
    sz = rPr.find(qn('w:sz'))
    if sz is None:
        sz = OxmlElement('w:sz')
        rPr.append(sz)
    sz.set(qn('w:val'), str(size * 2))
    szCs = rPr.find(qn('w:szCs'))
    if szCs is None:
        szCs = OxmlElement('w:szCs')
        rPr.append(szCs)
    szCs.set(qn('w:val'), str(size * 2))
    run.bold = bold
    run.font.color.rgb = RGBColor(*color)


def _rtl(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    bidi.set(qn('w:val'), '1')
    pPr.append(bidi)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def add_title(doc, text, size, color, after=4):
    p = doc.add_paragraph()
    _rtl(p)
    r = p.add_run(text)
    _arabic_font(r, size, bold=True, color=color)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    return p


def add_para(doc, parts, size=10, after=4, line_spacing=1.2):
    p = doc.add_paragraph()
    _rtl(p)
    if isinstance(parts, str):
        parts = [(parts, False)]
    for text, bold in parts:
        r = p.add_run(text)
        _arabic_font(r, size, bold=bold, color=TEXT)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line_spacing
    return p


def shade_cell(cell, hex_fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), hex_fill)
    tcPr.append(shd)


def set_table_rtl(table):
    tblPr = table._tbl.tblPr
    bidi = OxmlElement('w:bidiVisual')
    bidi.set(qn('w:val'), '1')
    tblPr.append(bidi)


def add_kv_table(doc, header_fill, rows, col_widths_cm, header_text=None,
                 row_size=9, bold_first_col=True):
    n_cols = len(rows[0])
    table = doc.add_table(rows=(1 if header_text else 0) + len(rows), cols=n_cols)
    set_table_rtl(table)
    table.autofit = False

    if header_text:
        for i, ht in enumerate(header_text):
            cell = table.rows[0].cells[i]
            cell.text = ''
            p = cell.paragraphs[0]
            _rtl(p)
            r = p.add_run(ht)
            _arabic_font(r, row_size + 1, bold=True, color=(255, 255, 255))
            shade_cell(cell, header_fill)
            cell.width = Cm(col_widths_cm[i])
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)

    start = 1 if header_text else 0
    for ri, row in enumerate(rows, start=start):
        for ci, val in enumerate(row):
            cell = table.rows[ri].cells[ci]
            cell.text = ''
            p = cell.paragraphs[0]
            _rtl(p)
            r = p.add_run(val)
            _arabic_font(r, row_size, bold=(ci == 0 and bold_first_col),
                         color=PRIMARY if (ci == 0 and bold_first_col) else TEXT)
            cell.width = Cm(col_widths_cm[ci])
            if (ri - start) % 2 == 1:
                shade_cell(cell, 'FAF6EE')
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.line_spacing = 1.15

    return table


def add_callout(doc, text, fill, accent_hex):
    table = doc.add_table(rows=1, cols=1)
    set_table_rtl(table)
    cell = table.rows[0].cells[0]
    cell.text = ''
    p = cell.paragraphs[0]
    _rtl(p)
    if isinstance(text, str):
        r = p.add_run(text)
        _arabic_font(r, 10, bold=False, color=TEXT)
    else:
        for t, bold in text:
            r = p.add_run(t)
            _arabic_font(r, 10, bold=bold, color=TEXT)
    shade_cell(cell, fill)
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), '24')
    left.set(qn('w:color'), accent_hex)
    borders.append(left)
    tcPr.append(borders)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.3
    cell.width = Cm(17)
    return table


def add_compact_gap_row(doc, num, title, body, accent_hex='C39B5C'):
    table = doc.add_table(rows=1, cols=2)
    set_table_rtl(table)
    table.autofit = False

    cell_l = table.rows[0].cells[0]
    cell_l.text = ''
    p_l = cell_l.paragraphs[0]
    _rtl(p_l)
    r_num = p_l.add_run(num + '  ')
    _arabic_font(r_num, 18, bold=True, color=ACCENT)
    r_title = p_l.add_run(title)
    _arabic_font(r_title, 10, bold=True, color=PRIMARY)
    cell_l.width = Cm(4.5)
    shade_cell(cell_l, 'FAF6EE')
    p_l.paragraph_format.space_before = Pt(2)
    p_l.paragraph_format.space_after = Pt(2)
    p_l.paragraph_format.line_spacing = 1.15

    cell_r = table.rows[0].cells[1]
    cell_r.text = ''
    p_r = cell_r.paragraphs[0]
    _rtl(p_r)
    r_body = p_r.add_run(body)
    _arabic_font(r_body, 9, color=TEXT)
    cell_r.width = Cm(12.5)
    p_r.paragraph_format.space_before = Pt(2)
    p_r.paragraph_format.space_after = Pt(2)
    p_r.paragraph_format.line_spacing = 1.3

    tcPr = cell_l._tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), '20')
    left.set(qn('w:color'), accent_hex)
    borders.append(left)
    tcPr.append(borders)


def divider(doc, before=4, after=4):
    p = doc.add_paragraph()
    _rtl(p)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:color'), 'C39B5C')
    pBdr.append(bottom)
    pPr.append(pBdr)


# =====================================================
# Document
# =====================================================

doc = Document()

for s in doc.sections:
    s.left_margin = Cm(1.8)
    s.right_margin = Cm(1.8)
    s.top_margin = Cm(1.5)
    s.bottom_margin = Cm(1.5)

# ── Header band ─────────────────────
title_p = doc.add_paragraph()
_rtl(title_p)
r = title_p.add_run('مَركز السيرة النبوية الجامع')
_arabic_font(r, 20, bold=True, color=PRIMARY)
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_p.paragraph_format.space_after = Pt(0)

sub_p = doc.add_paragraph()
_rtl(sub_p)
r = sub_p.add_run('تشخيص استراتيجي على أوسع مَسح عالمي لحقل السيرة المؤسسي')
_arabic_font(r, 11, bold=False, color=ACCENT)
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_p.paragraph_format.space_after = Pt(2)

meta_p = doc.add_paragraph()
_rtl(meta_p)
r = meta_p.add_run('٤٧٢ كياناً  ·  ٥٢ دولة  ·  ٢٠ جولة بحثية موثَّقة  ·  مايو ٢٠٢٦')
_arabic_font(r, 9, bold=False, color=MUTED)
meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta_p.paragraph_format.space_after = Pt(4)
divider(doc, before=0, after=6)

# ===========================================================================
# الصفحة ١
# ===========================================================================

add_title(doc, '١. لا يوجد كيان مماثل في العالم', 14, PRIMARY, after=4)

add_callout(doc, [
    ('بعد مَسح ', False),
    ('٤٧٢ كياناً ', True),
    ('في ', False),
    ('٥٢ دولة ', True),
    ('عبر ', False),
    ('٢٠ جولة بحثية موثَّقة (٢٠٢٥–٢٠٢٦)', True),
    ('، لا يوجد كيان واحد يَجمع الشمول والتكامل الذي تَقترحه الوثيقة: ', False),
    ('المنصّة الجامعة + الموسوعة الموضوعية الأخلاقية + الجامع للنصوص + الأطلس التفاعلي + قسم الطفل + القسم الإنجليزي + القسم البحثي + شبهات وردود', True),
    (' — كلّها في كيان واحد لنفس الجمهور.', False),
], fill='FEF3C7', accent_hex='D97706')

add_para(doc, '', size=6, after=2)
add_title(doc, 'سعة المسح ودقّته', 11, PRIMARY, after=3)

scope_rows = [
    ['الكيانات المرصودة', '٤٧٢ (٤٣٦ جوهري + ٣٤ ثانوي + ٢ فجوة موثَّقة)'],
    ['الدول المغطّاة', '٥٢ دولة عبر ٧ أقاليم'],
    ['الجولات البحثية', '٢٠ جولة موثَّقة، منها ٣ موجات استقراء شامل'],
    ['متوسّط جودة الكيان', '٨٦.٦ / ١٠٠ — صفر كيان تحت ٦٠'],
    ['نسبة التحقُّق المكتبي', '٩٢.٢٪ بزيارة الموقع الرسمي + المصدر'],
    ['حقول البطاقة', '٣٥ حقلًا منضبطة بمخطّط JSON Schema'],
    ['حقول مكتملة بـ١٠٠٪', '١٤ حقلًا (هوية، نوع، إقليم، تخصص، منهجية…)'],
    ['التدقيق الآلي', 'مدقّق آلي يَمنع أي خطأ قبل النشر — pre-commit + CI'],
    ['الترخيص', 'مفتوح المصدر CC BY 4.0 — قابل للتحقّق من أي طرف'],
]
add_kv_table(doc, '0A4D68', scope_rows, [5, 12],
             header_text=['المؤشّر', 'القيمة'], row_size=9)

add_para(doc, '', size=6, after=2)
add_title(doc, 'أقرب النظراء — ولا يَكتمل أيّهم', 11, PRIMARY, after=3)

peers_rows = [
    ['وقف السيرة التركي (Siyer Vakfı)', 'منصّة قوية بالتركية — بلا قسم موضوعي أخلاقي'],
    ['دار الرسول الأعظم العراقية', 'شيعي إمامي — يَفتقد البُعد الأمميّ السنّي'],
    ['الهيئة العالمية للتعريف بالرسول (الرابطة)', 'تَعريف بـ٨ لغات — بلا موسوعية موضوعية تطبيقية'],
    ['لجنة الموسوعة المصرية للسنة (٢٠٢٤)', 'ناشئة — تَركز على السنة لا السيرة الموضوعية'],
    ['موسوعة العاملي «الصحيح من سيرة النبي الأعظم» (٣٥ مجلد)', 'شيعي تأريخي روائي — لا موضوعي أخلاقي'],
]
add_kv_table(doc, '0A4D68', peers_rows, [6.5, 10.5],
             header_text=['الكيان النظير', 'الفجوة التي يَتركها'], row_size=9)

doc.add_page_break()

# ===========================================================================
# الصفحة ٢
# ===========================================================================

add_title(doc, '٢. ما لا يَخدمه أحد — خمس فجوات بنيوية', 14, PRIMARY, after=4)

add_para(doc, [
    ('خمس فجوات موثَّقة بالبيانات، كلٌّ منها تَستحق منتجاً قائماً بذاته. أيّ مشروع يَطمح للأمميّة يَنبغي أن يُموضع نفسه على هذه الفجوات تَحديداً، لا على ', False),
    ('منصّة جامعة عامة', True),
    ('.', False),
], size=10, after=6)

add_compact_gap_row(
    doc, '١', 'الموسوعة الموضوعية الأخلاقية',
    'لا يوجد كيان يُقدّم موسوعة موضوعية للسيرة في الأخلاق التطبيقية: الإدارة النبوية للزواج، فرق العمل، التربية، الأزمات، التواصل، الأخلاق الاقتصادية، الدبلوماسية. أقرب الموجود (سلسلة بنت الشاطئ، ٥ مجلدات) جزئيّ. بشهادة البيانات: أكبر إضافة مَمكنة في الحقل عالمياً.'
)
add_para(doc, '', size=4, after=2)

add_compact_gap_row(
    doc, '٢', 'السيرة من كامل السنّة',
    'جمع نصوص السيرة من جميع كتب السنّة (الجوامع والصحاح والسنن والمسانيد والأجزاء) ثم تَصنيفها زمنياً ومَوضوعياً. منهج لم يُنفَّذ بهذا الشمول في الـ٤٧٢ كياناً. أقرب نظير: موسوعة العاملي الشيعية بمنهج روائي مختلف.'
)
add_para(doc, '', size=4, after=2)

add_compact_gap_row(
    doc, '٣', 'الأطلس التفاعلي ثلاثي الأبعاد',
    'الموجود: الأطلس الدارة (PDF ساكن)، SirahMaps (محدود بمكة)، ديوراما إسطنبول (فيزيائي). التفاعلية الرقمية ثلاثية الأبعاد للسيرة (بيت خديجة ← غار حراء ← داخل الغار ← نزول الوحي) لا نظير لها.'
)
add_para(doc, '', size=4, after=2)

add_compact_gap_row(
    doc, '٤', 'قناة سيرة تلفزيونية مَكرَّسة',
    'بعد فحص ٤٧٢ كياناً: صفر قناة عالمياً مَكرَّسة كلّياً للسيرة. الأقرب: قناة السنة السعودية (سيرة جزئية). فجوة سوقية مُؤكَّدة، خاصة في صيغة OTT (الإنترنت) بدل البثّ التقليدي.'
)
add_para(doc, '', size=4, after=2)

add_compact_gap_row(
    doc, '٥', 'منظومة عربية–إنجليزية مُتكاملة',
    'الفضاء الإنجليزي قوي (Yaqeen, AlMaghrib, Cambridge Muslim College) لكنّه مستقل تماماً عن العربي. الترجمة لا تَكفي — يَلزم منهج مُصاغ من البداية ثنائي اللسان.'
)

doc.add_page_break()

# ===========================================================================
# الصفحة ٣
# ===========================================================================

add_title(doc, '٣. ١٥ شراكة استراتيجية على ثلاث طبقات', 14, PRIMARY, after=4)

add_para(doc, [
    ('البيانات تَكشف ', False),
    ('١٥ كياناً ', True),
    ('يَستحق التعامل المُمنهج، مُصنَّفة في ثلاث فئات بحسب طبيعة التعاون.', False),
], size=10, after=6)

add_title(doc, 'فئة أ — شراكات تَكاملية (مشاركة في البنية الأكاديمية)', 10, (21, 128, 61), after=2)
cat_a = [
    ['وقف السيرة التركي', 'تركيا', 'تَبادل المنهجية، نموذج الحوكمة، ترجمة الإنتاج'],
    ['كرسي ابن باز للسيرة (الجامعة الإسلامية بالمدينة)', 'السعودية', 'اعتماد أكاديمي + إشراف بحثي على الإصدارات'],
    ['لجنة الموسوعة المصرية للسنة (المجلس الأعلى)', 'مصر', 'تَوزيع أدوار: هم السنة، نحن السيرة الموضوعية'],
    ['مجمع الملك سلمان للحديث النبوي (المدينة)', 'السعودية', 'تَحقيق نصوص السيرة من كامل السنة'],
    ['دار المصطفى تريم (الحبيب عمر بن حفيظ)', 'اليمن', 'قسم الحديث الصوفي + نموذج التأثير المعنوي'],
]
add_kv_table(doc, '15803D', cat_a, [6, 1.6, 9.4],
             header_text=['الكيان', 'البلد', 'طبيعة الشراكة'], row_size=8.5)
add_para(doc, '', size=4, after=2)

add_title(doc, 'فئة ب — شراكات إنتاج (محتوى مشترك)', 10, (180, 134, 11), after=2)
cat_b = [
    ['Yaqeen Institute', 'الولايات المتحدة', 'محتوى إنجليزي مشترك + شهادات قَصيرة'],
    ['AlMaghrib Institute', 'كندا/المملكة المتحدة', 'دورات إنجليزية معتمدة بالمنهجية'],
    ['لجنة السيرة الوطنية البنغلاديشية', 'بنغلاديش', 'الترجمة البنغالية + شراكة أمميّة لـ١٥٣ مليون مسلم'],
    ['Diyanet TV', 'تركيا', 'إنتاج تلفزيوني بالعربية والتركية'],
    ['كراسي الطب النبوي السعودية (KAU + PSAU)', 'السعودية', 'إنتاج مُشترك في تَخصُّص فرعي'],
]
add_kv_table(doc, 'B4860B', cat_b, [6, 2.5, 8.5],
             header_text=['الكيان', 'البلد', 'طبيعة الشراكة'], row_size=8.5)
add_para(doc, '', size=4, after=2)

add_title(doc, 'فئة ج — شراكات مرجعية (الاستئناس بالتجربة)', 10, (120, 53, 15), after=2)
cat_c = [
    ['مركز ابن القطان لتحقيق التراث', 'المغرب', 'نموذج التحقيق الكلاسيكي'],
    ['المعرض والمتحف الدولي للسيرة (الرابطة - المدينة)', 'السعودية', 'نموذج المتحف التفاعلي'],
    ['مهرجان مولد لامو السنوي', 'كينيا', 'نموذج التفاعل الجماهيري الإفريقي'],
    ['مجلة IJICM الإندونيسية', 'إندونيسيا', 'نموذج المجلة المحكَّمة في تَخصُّص فرعي'],
    ['IslamHouse + Dorar.net (قسم السيرة)', 'السعودية', 'نموذج الأرشيف الرقمي متعدد اللغات'],
]
add_kv_table(doc, '78350F', cat_c, [6, 2.5, 8.5],
             header_text=['الكيان', 'البلد', 'الدرس المرجعي'], row_size=8.5)
add_para(doc, '', size=4, after=4)

add_callout(doc, [
    ('الخلاصة: ', True),
    ('المشروع يَدخل حقلاً نشطاً لكنّه ليس مُكتظّاً في النقاط الجوهرية. الإضافة الأقوى = الموضوعية الأخلاقية التطبيقية + الجمع المنهجي من كامل السنة. التوقيت حرج — اللجنة المصرية ستُكمل موسوعتها خلال ٣–٥ سنوات. ', False),
    ('التَخصُّص قبل الانتشار، الشراكة قبل المنافسة، الإطلاق قبل الكَمال.', True),
], fill='FEE2E2', accent_hex='B91C1C')

# Save
out = '/Users/mhmd/Documents/seerah-encyclopedia/مشروع_السيرة_تشخيص_استراتيجي.docx'
doc.save(out)
print(f"✓ Saved: {out}")
import os
print(f"  Size: {os.path.getsize(out):,} bytes")
